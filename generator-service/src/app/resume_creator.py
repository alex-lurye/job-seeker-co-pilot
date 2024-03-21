from docx import Document
from docx.shared import Pt
# Imports the Google Cloud client library
from google.cloud import storage
import re
from docx.text.paragraph import Paragraph
from docx.oxml.shared import OxmlElement
from copy import deepcopy

# Instantiates a client
#storage_client = storage.Client()

def insert_paragraph_after(paragraph, text=None, style=None):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style
    return new_para


def find_section_start(paragraphs: Paragraph, start_placeholder: str):
    """
    Find the index of the paragraph that contains the start of the section.
    """
    for i, paragraph in enumerate(paragraphs):
        if start_placeholder in paragraph.text:
            return i
    return None

def replace_placeholders_in_range(doc, start_index, end_index, replacements):
    """
    Replace placeholders in a range of paragraphs with actual data.
    """
    for i in range(start_index, end_index):
        paragraph = doc.paragraphs[i]
        for placeholder, actual in replacements.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(actual))

def replace_section_content(doc, section_start_index, section_length, replacements):
    """
    Replace the content of the original section using shared logic.
    """
    replace_placeholders_in_range(doc, section_start_index, section_start_index + section_length, replacements)

def clone_and_add_section(doc, start_index, section_length, replacements):
    
    end_index = start_index + section_length
    insert_after_para = doc.paragraphs[start_index + section_length - 1]
    # Clone section by adding new paragraphs at the end of the document
    for i in range(start_index, end_index):
        new_paragraph_text = doc.paragraphs[i].text
        for placeholder, actual in replacements.items():
            new_paragraph_text = new_paragraph_text.replace(placeholder, str(actual))
        
        insert_after_para = insert_paragraph_after(insert_after_para, new_paragraph_text)
    
    insert_paragraph_after(insert_after_para, "") # Add an empty paragraph at the end

def replace_placeholder_paragraph(paragraph, placeholder, new_content):
    """
    Replace placeholders in the paragraph with new content.
    """
    if placeholder in paragraph.text:
        inline = paragraph.runs
        for i in range(len(inline)):
            if placeholder in inline[i].text:
                text = inline[i].text.replace(placeholder, new_content)
                inline[i].text = text

def replace_placeholder(doc: Document, placeholder: str, new_content: str):
    """
    Replace placeholders in the document with new content.
    """
    for paragraph in doc.paragraphs:
        replace_placeholder_paragraph(paragraph, placeholder, new_content)


def clone_paragraph(paragraph, old_text, new_text):
    """
    Clone a paragraph, insert it after the original, and replace its text.
    """
    # Access the parent element of the paragraph (typically a _Body object)
    p_parent = paragraph._element.getparent()
    # Find the index of the paragraph to clone
    p_index = p_parent.index(paragraph._element)
    
    # Copy the paragraph element
    new_p_element = deepcopy(paragraph._element)
    # Insert the copied paragraph into the document
    p_parent.insert(p_index + 1, new_p_element)
    
    # Access the new paragraph through the Document object (for high-level operations like text replacement)
    new_paragraph = Paragraph(new_p_element, paragraph._parent)
    # Replace the old placeholder with the new text
    replace_placeholder_paragraph(new_paragraph, old_text, new_text)
    
    return new_paragraph

def create_resume(prof_summary: str, key_competiences: str, experiences: list, educations: list, skills: list):

    template_path = './resources/Resume_template.docx'
    
    doc = Document(template_path)
    
    replace_placeholder(doc, '{YOUR_PROFILE}', prof_summary)
    
    # Split the response text into lines
    lines = key_competiences.strip().split('\n')

    # Filter lines that start with a digit followed by a period to identify bullet points
    bullet_points = [line for line in lines if line.strip().startswith(('1.', '2.', '3.', '4.'))]
    
    clean_bullet_points = [re.sub(r'^\d+\.\s*', '', point).strip() for point in bullet_points]

    # yes I'm lazy
    replace_placeholder(doc,  '{YOUR_COMPETENCY_1}', clean_bullet_points[0])
    replace_placeholder(doc,  '{YOUR_COMPETENCY_2}', clean_bullet_points[1])
    replace_placeholder(doc,  '{YOUR_COMPETENCY_3}', clean_bullet_points[2])
    replace_placeholder(doc,  '{YOUR_COMPETENCY_4}', clean_bullet_points[3])
    
    replace_placeholder(doc, '{YOUR_SKILLS}', skills[0]['description'])

    # Find the start of the template experience section 
    start_placeholder = '{YOUR_EXPERIENCE_TITLE_1}'  
    section_start_index = find_section_start(doc.paragraphs, start_placeholder)
    experience_section_length = 3  # The number of paragraphs in the experience section
    if section_start_index is not None:
    
        def create_replacements_dict(experience):
            replacements = {}
            fields = {
                'positionTitle': '{YOUR_EXPERIENCE_TITLE_1}',
                'startYear': '{EXPERIENCE_START_YEAR_1}',
                'endYear': '{EXPERIENCE_END_YEAR_1}',
                'company': '{EXPERIENCE_COMPANY_NAME_1}',
                'countryIso2': '{EXPERIENCE_COUNTRY_1}',
                'description': '{EXPERIENCE_DESCRIPTION_1}',
            }
            for field, replacement_key in fields.items():
                if field in experience:
                    replacements[replacement_key] = experience[field]
            return replacements
   
        # Clone and add sections for all experiences except the first one
        for experience in experiences[1:]:
            replacements = create_replacements_dict(experience)
            clone_and_add_section(doc, section_start_index, experience_section_length, replacements)
        
        # Handle the first experience separately
        # Now the first experience's placeholders are replaced in the original section
        first_experience_replacements = create_replacements_dict(experiences[0])
        replace_section_content(doc, section_start_index, experience_section_length, first_experience_replacements)

    # Find the start of the template education section
    start_placeholder = '{YOUR_DEGREE_1}'
    section_start_index = find_section_start(doc.paragraphs, start_placeholder)
    education_section_length = 2  # The number of paragraphs in the education section
    
    if section_start_index is not None:
        def create_replacements_dict(education):
            replacements = {}
            fields = {
                'institution': '{YOUR_UNIVERSITY_NAME_1}',
                'startYear': '{DEGREE_START_YEAR_1}',
                'endYear': '{DEGREE_END_YEAR_1}',
                'fieldOfStudy': '{YOUR_DEGREE_1}',
            }
            for field, replacement_key in fields.items():
                if field in education:
                    replacements[replacement_key] = education[field]
            
            return replacements
        
        # Clone and add sections for all educations except the first one
        for education in educations[1:]:
            replacements = create_replacements_dict(education)
            clone_and_add_section(doc, section_start_index, education_section_length, replacements)
        
        # Handle the first education separately
        # Now the first education's placeholders are replaced in the original section
        first_education_replacements = create_replacements_dict(educations[0])
        replace_section_content(doc, section_start_index, education_section_length, first_education_replacements)
    
    doc.save('/tmp/resume.docx')

    return '/tmp/resume.docx'